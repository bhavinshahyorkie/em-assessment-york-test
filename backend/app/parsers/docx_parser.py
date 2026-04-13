"""DOCX / OOXML Word text extraction (`python-docx` plus ZIP/XML fallback for strict exports)."""

from __future__ import annotations

import io
import zipfile
import xml.etree.ElementTree as ET

import docx
from docx.document import Document as DocxDocument

from .base import DocumentParser, ParsedDocument


def _extract_docx_text(document: DocxDocument) -> str:
    """Paragraphs plus table cells — many resumes put content only in tables."""
    chunks: list[str] = []
    for p in document.paragraphs:
        t = p.text.strip()
        if t:
            chunks.append(t)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text.strip()
                    if t:
                        chunks.append(t)
    return "\n".join(chunks)


def _ooxml_text_from_docx_bytes(raw: bytes) -> str | None:
    """
    Second-pass extraction: read OOXML directly from the docx zip (Office Open XML).

    python-docx is strict; many real-world .docx files (Google Docs export, WPS, etc.) still
    validate as a zip with word/document.xml but fail Document(...). We collect all w:t text runs.
    """
    try:
        zf = zipfile.ZipFile(io.BytesIO(raw), "r")
    except zipfile.BadZipFile:
        return None

    if "word/document.xml" not in zf.namelist():
        return None

    parts_to_read: list[str] = ["word/document.xml"]
    for name in sorted(zf.namelist()):
        if (name.startswith("word/header") or name.startswith("word/footer")) and name.endswith(".xml"):
            parts_to_read.append(name)
    for extra in ("word/footnotes.xml", "word/endnotes.xml"):
        if extra in zf.namelist() and extra not in parts_to_read:
            parts_to_read.append(extra)

    seen: set[str] = set()
    ordered_paths: list[str] = []
    for p in parts_to_read:
        if p in zf.namelist() and p not in seen:
            seen.add(p)
            ordered_paths.append(p)

    chunks: list[str] = []
    for path in ordered_paths:
        try:
            with zf.open(path) as f:
                xml_bytes = f.read()
        except KeyError:
            continue
        chunks.append(_xml_collect_w_t_text(xml_bytes))

    merged = "\n".join(c for c in chunks if c and c.strip())
    return merged if merged.strip() else None


def _xml_collect_w_t_text(xml_bytes: bytes) -> str:
    """Pull text from WordprocessingML <w:t> runs (any namespace prefix)."""
    try:
        root = ET.fromstring(xml_bytes)
    except ET.ParseError:
        return ""

    out: list[str] = []
    for el in root.iter():
        if el.tag.endswith("}t"):
            if el.text:
                out.append(el.text)
            if el.tail:
                out.append(el.tail)
    return " ".join(out)


class DocxParser(DocumentParser):
    format_id = "docx"

    def supports(self, filename: str, content_type: str | None) -> bool:
        lower = filename.lower()
        if lower.endswith((".docx", ".docm")):
            return True
        ct = (content_type or "").lower().split(";")[0].strip()
        return ct in (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/vnd.ms-word.document.macroEnabled.12",
        )

    def parse(self, raw: bytes, source_label: str) -> ParsedDocument:
        text = ""
        primary_error: Exception | None = None

        try:
            document = docx.Document(io.BytesIO(raw))
            text = _extract_docx_text(document)
        except Exception as e:
            primary_error = e

        if not text.strip():
            fallback = _ooxml_text_from_docx_bytes(raw)
            if fallback and fallback.strip():
                text = fallback
            elif primary_error is not None:
                err_s = str(primary_error).lower()
                hint = (
                    " If the file is actually old Word .doc (binary), save as .docx or PDF from Word."
                    if "zip" in err_s
                    else ""
                )
                raise ValueError(
                    f"Could not read as DOCX ({source_label}): {primary_error}.{hint}"
                ) from primary_error
            else:
                raise ValueError(
                    f"DOCX contained no extractable text ({source_label}). "
                    "Try exporting again from Word/Google Docs, or use PDF."
                )

        return ParsedDocument(text=text.strip(), source_label=source_label, format_id=self.format_id)
